import streamlit as st
import pandas as pd
from io import BytesIO
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import PyPDF2
import re
from collections import Counter

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="CV Match AI",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- INITIALIZE SESSION STATE ---
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "The Job"
if 'job_role' not in st.session_state:
    st.session_state.job_role = ""
if 'cv_full_text' not in st.session_state:
    st.session_state.cv_full_text = ""
if 'job_desc' not in st.session_state:
    st.session_state.job_desc = ""
if 'input_method' not in st.session_state:
    st.session_state.input_method = "Upload File"
if 'manual_cv_data' not in st.session_state:
    st.session_state.manual_cv_data = {}
if 'follow_up_answers' not in st.session_state:
    st.session_state.follow_up_answers = {}
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {"matches": [], "missing": [], "score": 0}
if 'final_cv_data' not in st.session_state:
    st.session_state.final_cv_data = {}
if 'active_lang' not in st.session_state:
    st.session_state.active_lang = "English"

# --- CUSTOM CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=EB+Garamond:wght@400;700&family=Inter:wght@400;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', 'Inter', sans-serif;
    }

    :root {
        --brand-primary: #2563eb;
        --brand-dark: #1e293b;
        --brand-slate: #64748b;
        --bg-main: #f8fafc;
        --white: #ffffff;
        --divider: #e2e8f0;
    }

    .main { background-color: var(--bg-main); }

    .main-header {
        color: var(--brand-dark);
        font-weight: 800;
        font-size: 2.5rem;
        margin-bottom: 0.1rem;
        letter-spacing: -0.04em;
    }

    .sub-header {
        color: var(--brand-slate);
        font-size: 1.1rem;
        margin-bottom: 2rem;
        font-weight: 500;
    }

    .cv-preview-container {
        background-color: #f1f5f9;
        padding: 40px 20px;
        display: flex;
        justify-content: center;
    }

    .cv-page {
        background-color: white;
        width: 100%;
        max-width: 800px;
        padding: 50px 65px;
        box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.15);
        color: #111827;
        line-height: 1.5;
        font-family: 'Inter', sans-serif;
    }

    .cv-header-block { text-align: center; margin-bottom: 25px; }
    .cv-name { font-size: 2.25rem; font-weight: 800; color: #0f172a; margin-bottom: 10px; text-transform: uppercase; letter-spacing: 1px; }
    .cv-contact { font-size: 0.9rem; color: #475569; font-weight: 500; }
    .cv-header-divider { border: none; height: 2px; background-color: #1e293b; margin: 15px 0 25px 0; }
    .cv-section-title { font-size: 1rem; font-weight: 700; color: #1e293b; text-transform: uppercase; letter-spacing: 1.5px; margin-top: 25px; margin-bottom: 5px; border-bottom: 1px solid #cbd5e1; padding-bottom: 3px; }
    .cv-entry { margin-bottom: 15px; }
    .cv-entry-header { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 2px; }
    .cv-entry-title { font-weight: 700; font-size: 1.05rem; color: #0f172a; }
    .cv-entry-date { font-weight: 600; font-size: 0.85rem; color: #64748b; }
    .cv-entry-sub { font-weight: 600; font-size: 0.9rem; color: #475569; margin-bottom: 6px; }
    .cv-bullet-list { margin: 0; padding-left: 18px; }
    .cv-bullet { font-size: 0.92rem; margin-bottom: 4px; color: #334155; }
    .cv-skills-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 15px; margin-top: 10px; }
    .cv-skill-category { font-size: 0.9rem; }
    .cv-skill-label { font-weight: 700; color: #1e293b; margin-right: 5px; }

    .rtl { direction: rtl; text-align: right; }
    .rtl .cv-header-block { text-align: center; }
    .rtl .cv-entry-header { flex-direction: row-reverse; }
    .rtl .cv-bullet-list { padding-right: 18px; padding-left: 0; }

    .stTabs [data-baseweb="tab-list"] { gap: 8px; background-color: transparent; padding: 0; }
    .stTabs [data-baseweb="tab"] { height: 44px; white-space: pre; background-color: var(--white); border: 1px solid var(--slate-200); border-radius: 10px; padding: 0 20px; color: var(--slate-600); font-weight: 600; transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1); }
    .stTabs [data-baseweb="tab"]:hover { border-color: var(--brand-primary); color: var(--brand-primary); }
    .stTabs [aria-selected="true"] { background-color: var(--brand-primary) !important; color: var(--white) !important; border-color: var(--brand-primary) !important; box-shadow: 0 4px 12px rgba(37, 99, 235, 0.2); }

    .content-card { background-color: var(--white); padding: 2.5rem; border-radius: 20px; border: 1px solid rgba(226, 232, 240, 0.8); box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.04); margin-bottom: 2rem; }
    .section-label { font-weight: 700; color: var(--slate-900); margin-bottom: 0.5rem; display: block; font-size: 1.2rem; }

    .keyword-tag { display: inline-block; background-color: #eff6ff; color: #1e40af; padding: 4px 12px; border-radius: 9999px; margin: 4px; font-size: 0.85rem; font-weight: 600; border: 1px solid #dbeafe; }
    .stButton > button { border-radius: 12px; font-weight: 700; padding: 0.75rem 2rem; background-color: var(--white); border: 1px solid var(--slate-200); transition: all 0.2s ease; }
    .stButton > button:hover { border-color: var(--brand-primary); color: var(--brand-primary); transform: translateY(-2px); box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); }
    </style>
    """, unsafe_allow_html=True)

# --- CONSTANTS & DICTIONARIES ---

STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "if", "then", "else", "when", "at", "from", "by",
    "for", "with", "about", "against", "between", "into", "through", "during", "before",
    "after", "above", "below", "to", "up", "down", "in", "out", "on", "off", "over", "under",
    "again", "further", "once", "here", "there", "all", "any", "both", "each", "few", "more",
    "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than",
    "too", "very", "s", "t", "can", "will", "just", "don", "should", "now", "i", "me", "my",
    "myself", "we", "our", "ours", "ourselves", "you", "your", "yours", "yourself",
    "yourselves", "he", "him", "his", "himself", "she", "her", "hers", "herself", "it", "its",
    "itself", "they", "them", "their", "theirs", "themselves", "what", "which", "who", "whom",
    "this", "that", "these", "those", "am", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "having", "do", "does", "did", "doing"
}

# Words that sound meaningful but are NOT job requirements
FILLER_WORDS = {
    # adjectives/adverbs that describe quality, not skills
    "outstanding", "excellent", "strong", "good", "great", "best", "better", "new", "high",
    "low", "large", "small", "fast", "quick", "key", "main", "major", "minor", "various",
    "different", "multiple", "specific", "general", "overall", "effective", "efficient",
    "innovative", "creative", "dynamic", "motivated", "passionate", "dedicated", "driven",
    # generic business nouns that are not skills
    "ability", "skills", "knowledge", "understanding", "background", "foundation", "basis",
    "approach", "process", "solution", "solutions", "result", "results", "outcome", "impact",
    "focus", "area", "areas", "aspect", "aspects", "level", "levels", "type", "types",
    "way", "ways", "part", "parts", "point", "points", "set", "use", "used", "using",
    "provide", "support", "help", "ensure", "maintain", "manage", "develop", "build",
    "create", "design", "implement", "improve", "increase", "reduce", "deliver", "drive",
    # company/job description boilerplate
    "role", "position", "job", "candidate", "team", "company", "organization", "department",
    "environment", "culture", "opportunity", "responsibilities", "requirement", "requirements",
    "preferred", "required", "needed", "plus", "bonus", "ideal", "minimum", "maximum",
    "years", "year", "experience", "familiarity", "exposure", "proficiency", "ability",
    "degree", "bachelor", "master", "phd", "field", "fields", "related", "relevant",
    # words that caused bugs (from screenshots)
    "suppliers", "charge", "offline", "products", "insights", "cross", "time", "outstanding",
    "solutions", "teams", "oems", "infra", "bring", "compile", "image", "ongoing", "issues",
    "execute", "continuous", "improvement", "monitor", "product", "algo", "adas", "sectional",
    "charge", "deliver", "based", "order", "revolution", "bright", "hands", "passion",
    "difference", "lead", "come", "make", "prevent", "accident", "semi", "fully", "autonomous",
    "vehicle", "change", "drive", "joining", "department", "rd", "cross", "sectional",
}

SKILL_KEYWORDS = {
    "python", "sql", "java", "javascript", "react", "node", "html", "css", "git", "github",
    "r", "matlab", "tableau", "powerbi", "power bi", "excel", "bigquery", "pandas", "numpy",
    "scikit-learn", "tensorflow", "pytorch", "aws", "azure", "docker", "kubernetes", "unix",
    "linux", "bash", "spark", "hadoop", "c++", "c#", "php", "ruby", "swift", "kotlin",
    "typescript", "scala", "go", "rust", "mongodb", "postgresql", "mysql", "redis", "kafka",
    "airflow", "dbt", "looker", "databricks", "snowflake", "jira", "confluence", "jenkins",
}

# Maps job requirement phrases → what to look for in the CV
# This enables SEMANTIC matching: "analytical skills" in job → "statistics" in CV = MATCH
REQUIREMENT_PATTERNS = [
    # --- Technical Skills ---
    (r'\bpython\b', "python"),
    (r'\bsql\b', "sql"),
    (r'\br\b(?:\s+programming|\s+language)?', "r"),
    (r'\bmatlab\b', "matlab"),
    (r'\bexcel\b', "excel"),
    (r'\btableau\b', "tableau"),
    (r'\bpower\s*bi\b', "power bi"),
    (r'\bgit\b', "git"),
    (r'\blinux\b|\bunix\b', "linux"),
    (r'\bdocker\b', "docker"),
    (r'\baws\b', "aws"),
    (r'\bazure\b', "azure"),
    (r'\bmachine\s+learning\b|\bml\b', "machine learning"),
    (r'\bdeep\s+learning\b', "deep learning"),
    (r'\bstatistic', "statistics"),
    (r'\bdata\s+anal', "data analysis"),
    (r'\bdata\s+sci', "data science"),
    (r'\bdata\s+vis', "data visualization"),
    (r'\bdata\s+min', "data mining"),
    (r'\bregression\b', "regression"),
    (r'\bmodeling\b|\bmodelling\b', "modeling"),
    (r'\bvalidat', "validation"),
    (r'\btest\b|\btesting\b', "testing"),
    (r'\balgorithm', "algorithms"),
    (r'\bautomati', "automation"),
    (r'\bdebugging?\b', "debugging"),
    (r'\bdeployment\b', "deployment"),
    (r'\boptimiz', "optimization"),
    # --- Education ---
    (r'\bb\.?sc\b|\bbachelor', "bachelor degree"),
    (r'\bm\.?sc\b|\bmaster', "master degree"),
    (r'\bindustrial\s+engineering\b', "industrial engineering"),
    (r'\belectrical\s+engineering\b', "electrical engineering"),
    (r'\bcomputer\s+science\b', "computer science"),
    (r'\bstatistics\b', "statistics degree"),
    (r'\beconomics\b', "economics"),
    (r'\bnatural\s+sciences?\b', "natural sciences"),
    (r'\bmathematic', "mathematics"),
    # --- Languages ---
    (r'\bhebrew\b', "hebrew"),
    (r'\benglish\b', "english"),
    (r'\barabic\b', "arabic"),
    # --- Soft Skills (only meaningful ones) ---
    (r'\banalytic', "analytical skills"),
    (r'\bcommunication\b', "communication"),
    (r'\bteamwork\b|\bteam\s+player\b|\bcollaborat', "teamwork"),
    (r'\btime\s+management\b', "time management"),
    (r'\bproblem[\s\-]solv', "problem solving"),
    (r'\bproject\s+management\b', "project management"),
    (r'\bpresent', "presentations"),
    (r'\bwriting\b|\bwritten\b', "writing"),
]

# Additional domain patterns for fields beyond data science
DOMAIN_PHRASES = [
    # Data / ML
    (r'\bfraud\s+detection\b', "fraud detection"),
    (r'\brisk\s+(?:management|assessment|modeli)', "risk management"),
    (r'\bnlp\b|\bnatural\s+language\s+processing\b', "NLP"),
    (r'\bcomputer\s+vision\b', "computer vision"),
    (r'\btime\s+series\b', "time series"),
    (r'\ba/?b\s+testing\b', "A/B testing"),
    (r'\bfeature\s+engineering\b', "feature engineering"),
    (r'\bpredictive\s+(?:model|analytic)', "predictive modeling"),
    (r'\bforecasting?\b', "forecasting"),
    (r'\bbig\s+data\b', "big data"),
    # Software Engineering
    (r'\bci/?cd\b|\bcontinuous\s+(?:integration|deployment)\b', "CI/CD"),
    (r'\brest(?:ful)?\s+api\b|\bapi\s+develop', "REST API"),
    (r'\bmicroservices?\b', "microservices"),
    (r'\bagile\b', "Agile"),
    (r'\bscrum\b', "Scrum"),
    (r'\bdevops\b', "DevOps"),
    (r'\bcybersecurity\b|\binformation\s+security\b', "cybersecurity"),
    # Design
    (r'\bfigma\b', "Figma"),
    (r'\bphotoshop\b', "Photoshop"),
    (r'\bux\b|\buser\s+experience\b', "UX design"),
    (r'\bwireframe\b', "wireframing"),
    (r'\bprototyp', "prototyping"),
    # Marketing / Growth
    (r'\bgoogle\s+analytics\b', "Google Analytics"),
    (r'\bseo\b', "SEO"),
    (r'\bsem\b|\bpaid\s+search\b|\bgoogle\s+ads\b', "SEM/Google Ads"),
    (r'\bcontent\s+(?:marketing|strategy|creation)\b', "content marketing"),
    (r'\bsocial\s+media\b', "social media"),
    (r'\bcopywriting\b', "copywriting"),
    (r'\bbranding?\b', "branding"),
    (r'\bsalesforce\b', "Salesforce"),
    (r'\bhubspot\b', "HubSpot"),
    (r'\bcrm\b', "CRM"),
    # Finance / Business
    (r'\bfinancial\s+(?:model|analys|report)', "financial modeling"),
    (r'\bvaluation\b', "valuation"),
    (r'\baccounting\b', "accounting"),
    (r'\bauditin', "auditing"),
    (r'\bbudgeting?\b', "budgeting"),
    (r'\bkpi\b', "KPIs"),
    (r'\bsupply\s+chain\b', "supply chain"),
    (r'\bprocurement\b', "procurement"),
    (r'\bcfa\b', "CFA"),
    (r'\bcpa\b', "CPA"),
    # Healthcare / Science
    (r'\bclinical\s+(?:trial|research)\b', "clinical research"),
    (r'\bregulatory\s+(?:affairs|compliance)\b', "regulatory compliance"),
    (r'\bpharmaceutical\b', "pharmaceutical"),
    (r'\bbiostatistic', "biostatistics"),
    (r'\bgdpr\b', "GDPR"),
    # Engineering
    (r'\bcad\b|\bautocad\b', "CAD"),
    (r'\bembedded\s+system', "embedded systems"),
    (r'\badas\b', "ADAS"),
    (r'\bsignal\s+processing\b', "signal processing"),
    # Education / Assessment
    (r'\bpsychometric\b', "psychometric"),
    (r'\bgmat\b', "GMAT"),
    (r'\bcurriculum\b', "curriculum development"),
]

# Semantic equivalences: if job needs X, CV having Y counts as a match
SEMANTIC_EQUIVALENCES = {
    "machine learning": ["machine learning", "ml", "sklearn", "scikit", "tensorflow", "pytorch", "neural", "deep learning", "ai"],
    "data analysis": ["data analysis", "analytics", "statistical", "pandas", "numpy", "excel", "sql", "r"],
    "statistics": ["statistics", "statistical", "regression", "modeling", "r", "spss", "quantitative"],
    "statistics degree": ["statistics", "data science", "mathematics", "quantitative"],
    "sql": ["sql", "database", "postgresql", "mysql", "queries", "bigquery", "sqlite"],
    "python": ["python", "pandas", "numpy", "scikit", "django", "flask", "scipy"],
    "visualization": ["visualization", "tableau", "powerbi", "matplotlib", "ggplot", "plotting", "charts"],
    "data visualization": ["visualization", "tableau", "powerbi", "matplotlib", "ggplot", "plotting"],
    "git": ["git", "github", "gitlab", "version control", "vcs"],
    "teamwork": ["team", "collaborated", "hackathon", "cross-functional", "together", "cooperation"],
    "communication": ["communication", "presented", "report", "documentation", "written", "verbal"],
    "problem solving": ["problem solving", "algorithm", "debug", "optimized", "solution", "hackathon"],
    "analytical skills": ["statistics", "data analysis", "modeling", "research", "analytical", "quantitative"],
    "validation": ["validation", "testing", "evaluation", "benchmarking", "assessment", "qa"],
    "testing": ["testing", "test", "validation", "evaluation", "benchmarking", "qa"],
    "modeling": ["modeling", "regression", "statistical", "machine learning", "predictive"],
    "algorithms": ["algorithms", "data structures", "computer science", "python", "programming"],
    "automation": ["automation", "scripts", "pipeline", "automated", "python", "bash"],
    "linux": ["linux", "unix", "bash", "terminal", "command line"],
    "r": ["r", "rstudio", "ggplot", "tidyverse", "statistical computing"],
    "excel": ["excel", "spreadsheet", "pivot", "vlookup", "sheets"],
    "natural sciences": ["statistics", "physics", "chemistry", "biology", "mathematics", "data science"],
    "industrial engineering": ["industrial engineering", "engineering", "operations research"],
    "hebrew": ["hebrew"],
    "english": ["english"],
    "debugging": ["debugging", "debug", "troubleshoot", "bug", "testing", "qa"],
    "deployment": ["deployment", "deploy", "production", "devops", "ci/cd", "docker"],
    "optimization": ["optimization", "optimize", "performance", "tuning", "efficiency"],
    "forecasting": ["forecasting", "forecast", "prediction", "time series", "arima"],
    "NLP": ["nlp", "natural language", "text processing", "bert", "transformers", "spacy", "nltk"],
    "computer vision": ["computer vision", "image processing", "opencv", "cnn", "object detection"],
    "A/B testing": ["a/b testing", "ab testing", "experimentation", "hypothesis", "statistical testing"],
    "Agile": ["agile", "scrum", "sprint", "kanban", "jira"],
    "DevOps": ["devops", "ci/cd", "docker", "kubernetes", "jenkins", "pipeline"],
    "SEO": ["seo", "search engine", "keywords", "organic", "ranking"],
    "branding": ["branding", "brand", "identity", "design", "marketing"],
    "financial modeling": ["financial modeling", "excel", "valuation", "dcf", "forecast"],
    "accounting": ["accounting", "bookkeeping", "financial statements", "gaap"],
    "compliance": ["compliance", "regulatory", "gdpr", "policy", "audit"],
    "psychometric": ["psychometric", "assessment", "testing", "evaluation"],
}

GENERIC_TERMS = {
    "ability", "advanced", "bachelor", "common", "core", "critical", "degree", "entry",
    "familiarity", "field", "fields", "foundation", "job", "knowledge", "large", "like",
    "strong", "excellent", "good", "responsible", "requirement", "requirements",
    "responsibilities", "candidate", "company", "team", "work", "working", "looking",
    "opportunity", "motivated", "passion", "years", "experience", "role", "position",
    "plus", "must", "needed", "required", "preferred"
}

# Used by extract_requirements to filter dynamic word extraction
GENERIC_IGNORE_WORDS = {
    "experience", "ability", "strong", "skills", "knowledge", "background",
    "candidate", "responsibilities", "requirements", "preferred", "needed",
    "required", "position", "company", "opportunity", "motivated", "passionate",
    "results", "impact", "solution", "solutions", "provide", "support", "ensure",
    "manage", "develop", "build", "create", "deliver", "drive", "years", "degree",
    "field", "related", "relevant", "working", "environment", "organization",
    "excellent", "proficient", "proven", "demonstrated", "understanding",
    "familiarity", "exposure", "minimum", "maximum", "including", "following",
    "various", "multiple", "several", "please", "apply", "submit", "resume",
    "looking", "seeking", "hiring", "culture", "startup", "office", "flexible",
    "growing", "exciting", "innovative", "dynamic", "collaborative", "talented",
    "highly", "ideally", "preferably", "additionally", "responsibility",
    "qualification", "function", "purpose", "overview", "cover", "letter",
    "contact", "about", "bonus", "ideal", "within", "across", "other",
    "basis", "manner", "format", "level", "areas", "types", "parts",
    "aspect", "value", "focus", "scope", "point", "using", "given",
    "large", "small", "quick", "major", "minor", "general", "overall",
    "effective", "efficient", "specific", "clear", "detail", "success",
    "works", "joining", "willing", "passion", "cross", "bring", "align",
    "assist", "report", "serve", "learn", "perform", "participate",
    "contribute", "coordinate", "communicate", "ensure", "obtain",
}

ACTION_VERBS = {
    "assisted": "Engineered", "helped": "Optimized", "did": "Developed",
    "worked on": "Spearheaded", "responsible for": "Directed", "learned": "Mastered",
    "managed": "Architected", "wrote": "Formulated"
}

HEADINGS_MAP = {
    "English": ["Personal Details", "Professional Summary", "Education", "Professional Experience", "Projects",
                "Courses & Training", "Volunteering / Community", "Languages", "Skills", "Additional Information"],
    "Hebrew": ["פרטים אישיים", "תמצית מקצועית", "השכלה", "ניסיון תעסוקתי", "פרויקטים", "קורסים והכשרות",
               "התנדבות / קהילה", "שפות", "כישורים", "מידע נוסף"],
    "Arabic": ["البيانات الشخصية", "الملخص المهني", "التعليم", "الخبرة العملية", "المشاريع", "الدورات والتدريب",
               "التطوع / المجتمع", "اللغات", "المهارات", "معلومات إضافية"]
}

CV_TITLES = {"English": "Curriculum Vitae", "Hebrew": "קורות חיים", "Arabic": "السيرة الذاتية"}

SECTION_REGEX = {
    "personal_details": r"(personal details|contact|contact info|פרטים אישיים|البيانات الشخصية)",
    "summary": r"(summary|profile|about me|professional summary|תמצית|פרופיל|תמצית מקצועית|الملخص|ملخص مهني)",
    "education": r"(education|academic|degrees|השכלה|לימודים|التعليم|المؤهلات العلمية)",
    "experience": r"(experience|work|employment|history|professional experience|ניסיון|ניסיון תעסוקתי|الخبرة|الخبرة العملية)",
    "projects": r"(projects|key projects|academic projects|פרויקטים|פרוייקטים|المشاريع)",
    "courses_training": r"(courses|training|certification|certifications|קורסים|הכשרות|הסמכות|الدورات|التدريب)",
    "volunteering": r"(volunteering|community|extracurricular|volunteer|התנדבות|פעילות קהילתית|التطوع)",
    "languages": r"(languages|שפות|اللغات)",
    "skills": r"(skills|competencies|technologies|technical skills|כישורים|מיומנויות|יכולות|المهارات)"
}

SAMPLE_CV_DATA = {
    "personal_details": "Jane Doe | jane.doe@university.edu | +1 555 0123 | LinkedIn: janedoe-profile",
    "summary": "Third-year Computer Science student with strong foundations in software engineering and data analysis.",
    "education": "B.Sc. in Computer Science, State University (Expected May 2025)\nRelevant Coursework: Data Structures, Algorithms, Database Systems, Web Development",
    "experience": "Summer Research Intern, AI Lab (2023)\n- Developed data preprocessing scripts in Python.\n- Conducted performance benchmarking for machine learning models.",
    "projects": "OpenSource Contributor - Improved UI components for a React-based dashboard project.\nWeather Insight App - Built a responsive weather tracking application using OpenWeather API.",
    "courses_training": "AWS Cloud Practitioner Essentials, Advanced SQL Certification (Coursera)",
    "volunteering": "Peer Tutor - Provided technical assistance to first-year CS students in Java and C++.",
    "languages": "English (Native), French (Conversational)",
    "skills": "Python, Java, SQL, React, Git, Docker, Problem Solving, Technical Writing"
}


# ============================================================
# CORE FUNCTIONS
# ============================================================

def detect_language(text):
    if re.search(r'[\u0590-\u05FF]', text): return "Hebrew"
    if re.search(r'[\u0600-\u06FF]', text): return "Arabic"
    return "English"


def clean_text_simple(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\u0590-\u05FF\u0600-\u06FF\s]', ' ', text)
    return text


def extract_keywords(text):
    """Original keyword extractor — used for CV text only."""
    cleaned = clean_text_simple(text)
    words = cleaned.split()
    return [w for w in words if w not in STOPWORDS and w not in GENERIC_TERMS and len(w) > 2]


def extract_requirements(job_text):
    """
    Extracts REAL requirements from any job in any field.
    Combines pattern matching for known skills with dynamic extraction
    for domain-specific terms the patterns don't cover.
    Ignores generic business words and anything under 5 characters.
    """
    job_lower = job_text.lower()
    requirements = {}

    # 1. Known tech / data science patterns
    for pattern, label in REQUIREMENT_PATTERNS:
        if re.search(pattern, job_lower):
            requirements[label] = True

    # 2. Broader domain patterns (marketing, finance, healthcare, etc.)
    for pattern, label in DOMAIN_PHRASES:
        if re.search(pattern, job_lower):
            requirements[label] = True

    # 3. Capitalised mid-sentence words → likely tool or product names
    # e.g. "experience with Figma and Sketch" → extracts "Figma", "Sketch"
    for m in re.finditer(r'(?<=[a-z0-9,;:])\s+([A-Z][a-zA-Z0-9\+\#\.]{2,})\b', job_text):
        token = m.group(1)
        t = token.lower()
        if t not in GENERIC_IGNORE_WORDS and t not in STOPWORDS:
            requirements[t] = True

    # 4. Words in bullet / numbered requirement lines (5+ chars, not ignored)
    for line in job_text.split('\n'):
        stripped = line.strip()
        is_bullet = stripped.startswith(('-', '•', '*', '·', '–', '○', '▪'))
        is_numbered = bool(re.match(r'^\d+[\.\)]\s', stripped))
        if is_bullet or is_numbered:
            for word in re.findall(r'\b([a-zA-Z][a-zA-Z0-9\-]{4,})\b', stripped):
                w = word.lower()
                if w not in GENERIC_IGNORE_WORDS and w not in STOPWORDS and w not in requirements:
                    requirements[w] = True

    return list(requirements.keys())


def smart_semantic_match(cv_text, job_requirements):
    """
    Checks each job requirement against the CV text semantically.
    Uses SEMANTIC_EQUIVALENCES so 'statistics degree' in job matches
    'Statistics and Data Science' in CV, even with different wording.
    Returns (matches, missing) as lists of requirement labels.
    """
    cv_lower = cv_text.lower()
    matches = []
    missing = []

    for req in job_requirements:
        found = False

        # Direct check: is the requirement word itself in the CV?
        if req in cv_lower:
            found = True

        # Semantic check: does the CV contain any equivalent term?
        if not found and req in SEMANTIC_EQUIVALENCES:
            for equiv in SEMANTIC_EQUIVALENCES[req]:
                if equiv in cv_lower:
                    found = True
                    break

        # Partial match: individual words of multi-word requirement
        if not found and ' ' in req:
            parts = req.split()
            if any(part in cv_lower for part in parts if len(part) > 3):
                found = True

        if found:
            matches.append(req)
        else:
            missing.append(req)

    return matches, missing


def calculate_score(matches, total_requirements):
    """
    Realistic scoring:
      raw  = (matched / total) * 100
      penalty = floor(missing / 3) * 5   (−5% per 3 missing requirements)
      score = clamp(raw − penalty, 15, 88)
    Cap prevents unrealistic 95%+ scores; floor keeps output meaningful.
    """
    m_count = len(matches) if hasattr(matches, "__len__") else int(matches)
    t_count = len(total_requirements) if hasattr(total_requirements, "__len__") else int(total_requirements)
    if t_count <= 0:
        return 0
    missing_count = t_count - m_count
    raw = (m_count / t_count) * 100
    penalty = (missing_count // 3) * 5
    score = raw - penalty
    return round(max(15, min(88, score)))


def validate_job_description(job_text):
    """
    Returns (is_valid, error_message).
    Rejects descriptions that are too short or obvious gibberish.
    """
    words = job_text.strip().split()
    if len(words) < 30:
        return False, "Please paste a real job description with responsibilities and requirements"
    unique_words = {w.lower() for w in words if len(w) > 2}
    if len(unique_words) < 10:
        return False, "Please paste a real job description with responsibilities and requirements"
    gibberish = {"test", "asdf", "nothing", "bla", "blah", "lorem", "ipsum", "aaa", "xxx", "qqq"}
    first_words = {w.lower() for w in words[:15]}
    if first_words & gibberish and len(unique_words) < 20:
        return False, "Please paste a real job description with responsibilities and requirements"
    return True, ""


def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except:
        return ""


def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except:
        return ""


def classify_sections_refined(text):
    sections = {
        "personal_details": "", "summary": "", "education": "", "experience": "",
        "projects": "", "courses_training": "", "volunteering": "", "languages": "",
        "skills": "", "additional_information": ""
    }
    current_section = "personal_details"
    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        low_line = clean_line.lower()
        found_header = False
        if len(clean_line) < 35:
            for key, pattern in SECTION_REGEX.items():
                if re.search(pattern, low_line):
                    current_section = key
                    found_header = True
                    break
        if not found_header:
            if current_section in ["experience", "education", "personal_details"] and len(clean_line) < 60:
                line_words = set(re.findall(r'\w+', low_line))
                if line_words.intersection(SKILL_KEYWORDS) and len(line_words) < 8:
                    sections["skills"] += clean_line + ", "
                    continue
            sections[current_section] += clean_line + "\n"
    for k in sections:
        sections[k] = re.sub(r'\n{2,}', '\n', sections[k].strip())
        if k == "skills":
            sections[k] = sections[k].strip(", ")
    return sections


def format_contact_info(text):
    lines = text.split('\n')
    if not lines: return "", ""
    name = lines[0].strip()
    rest = " | ".join([l.strip() for l in lines[1:] if l.strip()])
    return name, rest


def get_ai_suggestion(role, field):
    suggestions = {
        "summary": f"Ambitious {role} student focused on leveraging technical skills to build innovative solutions.",
        "experience": f"Collaborated with cross-functional teams to deliver {role}-related projects.",
        "projects": f"Developed a scalable application that demonstrates core {role} principles."
    }
    return suggestions.get(field, "Start typing your experience here...")


def group_skills_segregated(skills_text):
    if not skills_text: return {"PROFESSIONAL": [], "TECHNICAL": []}
    skills = [s.strip() for s in re.split(r'[,|\n]', skills_text) if s.strip()]
    tech_keywords = set(
        list(SKILL_KEYWORDS) + ["sql", "python", "aws", "docker", "git", "react", "tableau", "powerbi", "excel",
                                "data analysis"])
    technical, professional = [], []
    for s in skills:
        sl = s.lower()
        if any(kw in sl for kw in tech_keywords) or len(s) < 15:
            technical.append(s)
        else:
            professional.append(s)
    if not professional and technical:
        professional = [s for s in technical if len(s) > 10][:3]
        technical = [s for s in technical if s not in professional]
    return {"PROFESSIONAL": list(set(professional)), "TECHNICAL": list(set(technical))}


def integrate_refinement_smart(cv_data, answers, job_role, missing_keywords):
    upgraded_cv = cv_data.copy()

    def upgrade_text(text):
        if not text: return ""
        for weak, strong in ACTION_VERBS.items():
            text = re.sub(rf'\b{weak}\b', strong, text, flags=re.IGNORECASE)
        return text

    for sec in ["summary", "experience", "projects", "education"]:
        upgraded_cv[sec] = upgrade_text(upgraded_cv.get(sec, ""))
    answers_text = " ".join(answers.values()).lower()
    current_skills = upgraded_cv.get("skills", "").lower()
    new_skills = [mk for mk in missing_keywords if mk.lower() in answers_text and mk.lower() not in current_skills]
    if new_skills:
        upgraded_cv["skills"] = upgraded_cv.get("skills", "") + ", " + ", ".join(new_skills)
    achievement = ""
    for ans in answers.values():
        if re.search(r'\d+%|\d+x|improved|optimized|spearheaded', ans.lower()):
            achievement = ans.strip()
            break
    if achievement and upgraded_cv.get("experience"):
        exp_lines = upgraded_cv["experience"].split('\n')
        for i, line in enumerate(exp_lines):
            if line.strip().startswith(('-', '•', '*')):
                exp_lines[i] = f"* {achievement}"
                break
        upgraded_cv["experience"] = "\n".join(exp_lines)
    if job_role and upgraded_cv.get("summary"):
        if job_role.lower() not in upgraded_cv["summary"].lower():
            upgraded_cv[
                "summary"] = f"Goal-oriented professional with a focus on **{job_role}**. {upgraded_cv['summary']}"
    return upgraded_cv


def render_executive_cv_html(cv_data, lang="English"):
    is_rtl = lang in ["Hebrew", "Arabic"]
    l_class = "rtl" if is_rtl else "ltr"
    dir_attr = 'dir="rtl"' if is_rtl else 'dir="ltr"'
    h_map = HEADINGS_MAP[lang]
    name_contact = cv_data.get("personal_details", "CANDIDATE NAME | Contact Info")
    name, contact = format_contact_info(name_contact)
    if not name: name = "CANDIDATE NAME"
    html = f'<div class="cv-page {l_class}" {dir_attr}><div class="cv-header-block"><div class="cv-name">{name.upper()}</div><div class="cv-contact">{contact}</div><hr class="cv-header-divider"></div>'
    if cv_data.get("summary"):
        html += f'<div class="cv-section-title">{h_map[1].upper()}</div><div class="cv-entry" style="font-size:0.95rem; margin-top:8px;">{cv_data["summary"]}</div>'
    if cv_data.get("experience"):
        html += f'<div class="cv-section-title">{h_map[3].upper()}</div>'
        for entry in cv_data["experience"].split('\n\n'):
            lines = [l.strip() for l in entry.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-header"><span class="cv-entry-title">{lines[0]}</span></div><div class="cv-entry-sub">{lines[1] if len(lines) > 1 else ""}</div><ul class="cv-bullet-list">'
            for b in lines[2:]:
                style = 'style="list-style-type: circle; margin-left: 20px; font-size: 0.85rem; color: #475569;"' if b.startswith(
                    ('  ', '\t', '-')) else ''
                html += f'<li class="cv-bullet" {style}>{b.strip("- •*")}</li>'
            html += "</ul></div>"
    if cv_data.get("projects"):
        html += f'<div class="cv-section-title">{h_map[4].upper()}</div>'
        for proj in cv_data["projects"].split('\n\n'):
            lines = [l.strip() for l in proj.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-title">{lines[0]}</div><ul class="cv-bullet-list">'
            for b in lines[1:]:
                html += f'<li class="cv-bullet">{b.strip("- •*")}</li>'
            html += "</ul></div>"
    if cv_data.get("education"):
        html += f'<div class="cv-section-title">{h_map[2].upper()}</div>'
        for edu in cv_data["education"].split('\n\n'):
            lines = [l.strip() for l in edu.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-header"><span class="cv-entry-title">{lines[0]}</span></div><div class="cv-entry-sub">{lines[1] if len(lines) > 1 else ""}</div></div>'
    skills_data = group_skills_segregated(cv_data.get("skills", ""))
    if skills_data["PROFESSIONAL"]:
        html += f'<div class="cv-section-title">PROFESSIONAL SKILLS</div><div class="cv-entry" style="font-size:0.92rem; margin-top:8px;">{", ".join(skills_data["PROFESSIONAL"])}</div>'
    if skills_data["TECHNICAL"]:
        html += f'<div class="cv-section-title">TECHNICAL SKILLS</div><div class="cv-entry" style="font-size:0.92rem; margin-top:8px;">{", ".join(skills_data["TECHNICAL"])}</div>'
    html += "</div>"
    return html


def generate_docx_executive(cv_data, output_lang):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = docx.shared.Inches(1)
    is_rtl = output_lang in ["Hebrew", "Arabic"]
    h_map = HEADINGS_MAP[output_lang]
    style = doc.styles['Normal']
    style.font.name = 'Arial' if is_rtl else 'Calibri'
    style.font.size = Pt(11)
    name, contact = format_contact_info(cv_data.get("personal_details", ""))
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name.upper())
    run_name.bold = True
    run_name.font.size = Pt(22)
    p_contact = doc.add_paragraph(contact)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(6)
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    order = [("summary", 1), ("experience", 3), ("projects", 4), ("education", 2)]
    for key, h_idx in order:
        content = cv_data.get(key, "").strip()
        if content:
            h = doc.add_paragraph()
            if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_run = h.add_run(h_map[h_idx].upper())
            h_run.bold = True
            h_run.font.size = Pt(12)
            h.paragraph_format.space_before = Pt(12)
            if key in ["experience", "projects"]:
                for entry in content.split('\n\n'):
                    lines = [l.strip() for l in entry.split('\n') if l.strip()]
                    for i, line in enumerate(lines):
                        p = doc.add_paragraph()
                        if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if i == 0:
                            p.add_run(line).bold = True
                        elif i == 1 and key == "experience":
                            p.add_run(line).italic = True
                        else:
                            p_bullet = doc.add_paragraph(style='List Bullet')
                            if is_rtl: p_bullet.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            p_bullet.add_run(line.strip("- •*"))
            else:
                p = doc.add_paragraph(content)
                if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    skills_data = group_skills_segregated(cv_data.get("skills", ""))
    if skills_data["PROFESSIONAL"]:
        h = doc.add_paragraph()
        if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h.add_run("PROFESSIONAL SKILLS").bold = True
        doc.add_paragraph(", ".join(skills_data["PROFESSIONAL"]))
    if skills_data["TECHNICAL"]:
        h = doc.add_paragraph()
        if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h.add_run("TECHNICAL SKILLS").bold = True
        doc.add_paragraph(", ".join(skills_data["TECHNICAL"]))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# APP LAYOUT
# ============================================================

st.markdown('<div class="main-header">CV Match AI</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Premium Career Strategy & Tailoring Suite</div>', unsafe_allow_html=True)
tab_titles = ["The Job", "CV Content", "Analysis", "Refinement", "Preview", "Export"]
tabs = st.tabs(tab_titles)

with tabs[0]:
    st.markdown('<div class="content-card"><span class="section-label">Target Position & Requirements</span>',
                unsafe_allow_html=True)
    st.write("Define your target. Paste the description or insert a link to start the semantic analysis.")
    st.session_state.job_role = st.text_input("What is your target job title?", value=st.session_state.job_role,
                                              placeholder="e.g., Software Engineering Intern")
    input_mode = st.radio("How would you like to provide the job details?", ["Paste Job Text", "Insert Job Link"],
                          horizontal=True)
    if input_mode == "Insert Job Link":
        job_url = st.text_input("Job Posting URL:", placeholder="https://linkedin.com/jobs/...")
        if job_url:
            st.info("🔗 Link detected. (Simulated extraction in progress...)")
            if st.button("Extract from URL"):
                st.session_state.job_desc = f"Simulated extraction for: {st.session_state.job_role}. Requirements: Python, SQL, Git, Teamwork, Data Analysis."
    else:
        st.session_state.job_desc = st.text_area("Job Description:", value=st.session_state.job_desc, height=250,
                                                 placeholder="Paste requirements...")
    if st.button("✨ Load Demo Job Description"):
        st.session_state.job_role = "Junior Data Analyst"
        st.session_state.job_desc = "Requirements: Proficient in SQL and Python. Experience with Tableau or PowerBI. Knowledge of Git. Strong analytical skills. Bachelor degree in Statistics or Economics."
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[1]:
    st.markdown('<div class="content-card"><span class="section-label">Your Professional Profile</span>',
                unsafe_allow_html=True)
    st.session_state.input_method = st.radio("Choose input method:", ["Upload File", "Paste Text", "Manual Entry"],
                                             horizontal=True)
    if st.session_state.input_method == "Upload File":
        uploaded_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"])
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                st.session_state.cv_full_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                st.session_state.cv_full_text = extract_text_from_docx(uploaded_file)
            else:
                st.session_state.cv_full_text = str(uploaded_file.read(), "utf-8")
            st.success("File processed.")
    elif st.session_state.input_method == "Paste Text":
        st.session_state.cv_full_text = st.text_area("Paste CV Text:", value=st.session_state.cv_full_text, height=300)
        if st.button("✨ Auto-Structure Pasted Text"):
            st.session_state.manual_cv_data = classify_sections_refined(st.session_state.cv_full_text)
            st.success("Text structured successfully.")
    elif st.session_state.input_method == "Manual Entry":
        m_fields = {
            "summary": ("Your Hook / Career Goal", "Hardworking CS student..."),
            "experience": ("Work or Volunteer Experience", "Internships..."),
            "projects": ("Academic or Side Projects", "University projects..."),
            "education": ("Education", "University..."),
            "skills": ("Skills & Technologies", "Python..."),
            "personal_details": ("Full Name & Contact Info", "Name | Email...")
        }
        c1, c2 = st.columns(2)
        with c1:
            for key in ["personal_details", "summary", "education"]:
                st.write(f"**{m_fields[key][0]}**")
                if st.button(f"✨ Suggest {m_fields[key][0]}", key=f"ai_{key}"):
                    st.session_state.manual_cv_data[key] = get_ai_suggestion(st.session_state.job_role, key)
                st.session_state.manual_cv_data[key] = st.text_area(
                    m_fields[key][0], value=st.session_state.manual_cv_data.get(key, ""),
                    placeholder=m_fields[key][1], key=f"input_{key}",
                    label_visibility="collapsed", height=100)
        with c2:
            for key in ["experience", "projects", "skills"]:
                st.write(f"**{m_fields[key][0]}**")
                if st.button(f"✨ Suggest {m_fields[key][0]}", key=f"ai_{key}"):
                    st.session_state.manual_cv_data[key] = get_ai_suggestion(st.session_state.job_role, key)
                st.session_state.manual_cv_data[key] = st.text_area(
                    m_fields[key][0], value=st.session_state.manual_cv_data.get(key, ""),
                    placeholder=m_fields[key][1], key=f"input_{key}",
                    label_visibility="collapsed", height=100)
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[2]:
    cv_p = st.session_state.cv_full_text or any(st.session_state.manual_cv_data.values())
    if not st.session_state.job_desc or not cv_p:
        st.warning("Please provide Job Description and CV content.")
    else:
        is_valid, error_msg = validate_job_description(st.session_state.job_desc)
        if not is_valid:
            st.error(error_msg)
        else:
            st.markdown('<div class="content-card"><span class="section-label">Semantic Match Analysis</span>',
                        unsafe_allow_html=True)

            cv_text = st.session_state.cv_full_text if st.session_state.cv_full_text else "\n".join(
                st.session_state.manual_cv_data.values())

            job_requirements = extract_requirements(st.session_state.job_desc)
            matches, missing = smart_semantic_match(cv_text, job_requirements)
            score = calculate_score(matches, job_requirements)

            col1, col2 = st.columns([1, 2])
            with col1:
                st.write("### Match Score")
                score_color = "#16a34a" if score >= 70 else "#ea580c" if score >= 40 else "#dc2626"
                st.markdown(
                    f"<h1 style='color:{score_color}; font-size:4rem; margin:0;'>{score}%</h1>",
                    unsafe_allow_html=True)
                st.progress(score / 100)
                st.caption(f"{len(matches)} matched out of {len(job_requirements)} requirements")
            with col2:
                st.write("### Alignment")
                if score >= 70:
                    st.markdown(
                        "<div style='background:#dcfce7; color:#166534; padding:14px 18px; "
                        "border-radius:10px; font-weight:700; font-size:1.05rem;'>"
                        "Strong Match</div>",
                        unsafe_allow_html=True)
                elif score >= 40:
                    st.markdown(
                        "<div style='background:#ffedd5; color:#9a3412; padding:14px 18px; "
                        "border-radius:10px; font-weight:700; font-size:1.05rem;'>"
                        "Moderate Match</div>",
                        unsafe_allow_html=True)
                else:
                    st.markdown(
                        "<div style='background:#fee2e2; color:#991b1b; padding:14px 18px; "
                        "border-radius:10px; font-weight:700; font-size:1.05rem;'>"
                        "Weak Match — significant gaps detected</div>",
                        unsafe_allow_html=True)

            st.write("---")
            m_col, g_col = st.columns(2)
            with m_col:
                st.write(f"**Matched ({len(matches)}):**")
                if matches:
                    tags = "".join([
                        f'<span style="display:inline-block; background:#dcfce7; color:#166534; '
                        f'border:1px solid #bbf7d0; padding:4px 12px; border-radius:9999px; '
                        f'margin:3px; font-size:0.82rem; font-weight:600;">✓ {k}</span>'
                        for k in matches
                    ])
                    st.markdown(tags, unsafe_allow_html=True)
                else:
                    st.write("No matches found.")
            with g_col:
                st.write(f"**Missing ({len(missing)}):**")
                if missing:
                    tags = "".join([
                        f'<span style="display:inline-block; background:#fee2e2; color:#991b1b; '
                        f'border:1px solid #fecaca; padding:4px 12px; border-radius:9999px; '
                        f'margin:3px; font-size:0.82rem; font-weight:600;">✗ {k}</span>'
                        for k in missing
                    ])
                    st.markdown(tags, unsafe_allow_html=True)
                else:
                    st.success("No critical gaps found!")

            st.session_state.analysis_results = {"matches": matches, "missing": missing, "score": score}
            st.markdown('</div>', unsafe_allow_html=True)

with tabs[3]:
    st.markdown('<div class="content-card"><span class="section-label">Refinement Questions</span>',
                unsafe_allow_html=True)
    missing = st.session_state.analysis_results.get("missing", [])
    if not missing:
        st.write("No major gaps found!")
    else:
        q1 = f"The role requires **{missing[0]}**. Describe a project or experience where you used this."
        q2 = f"How have you demonstrated **{missing[1] if len(missing) > 1 else 'problem solving'}** in practice?"
        st.write(f"**Question 1:** {q1}")
        st.session_state.follow_up_answers["q1"] = st.text_area("A1",
                                                                value=st.session_state.follow_up_answers.get("q1", ""),
                                                                key="ans1", label_visibility="collapsed")
        st.write(f"**Question 2:** {q2}")
        st.session_state.follow_up_answers["q2"] = st.text_area("A2",
                                                                value=st.session_state.follow_up_answers.get("q2", ""),
                                                                key="ans2", label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[4]:
    st.markdown('<div class="content-card"><span class="section-label">Executive CV Preview</span>',
                unsafe_allow_html=True)
    st.session_state.active_lang = st.radio("Choose Document Language:", ["English", "Hebrew", "Arabic"],
                                            horizontal=True, key="lang_select")
    base_data = st.session_state.manual_cv_data.copy() if any(
        st.session_state.manual_cv_data.values()) else classify_sections_refined(st.session_state.cv_full_text)
    final_cv = integrate_refinement_smart(base_data, st.session_state.follow_up_answers, st.session_state.job_role,
                                          st.session_state.analysis_results.get("missing", []))
    st.session_state.final_cv_data = final_cv
    st.markdown('<div class="cv-preview-container">', unsafe_allow_html=True)
    st.markdown(render_executive_cv_html(final_cv, st.session_state.active_lang), unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    st.success("✨ Ready for download in the next tab.")
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[5]:
    st.markdown('<div class="content-card" style="text-align:center;"><span class="section-label">Final Export</span>',
                unsafe_allow_html=True)
    if not st.session_state.final_cv_data:
        st.warning("Please preview your CV first.")
    else:
        cv_f = st.session_state.final_cv_data
        u_name = cv_f["personal_details"].split('|')[0].strip().split()[0] if cv_f.get(
            "personal_details") else "Candidate"
        f_name = f"{u_name}_{st.session_state.job_role.replace(' ', '_')}_CV"
        c1, c2 = st.columns(2)
        with c1:
            txt_out = f"{f_name.upper()}\n{'=' * len(f_name)}\n"
            for k, v in cv_f.items():
                if v.strip(): txt_out += f"\n{k.upper()}\n{v}\n"
            st.download_button("📄 Download TXT", txt_out, file_name=f"{f_name}.txt", use_container_width=True)
        with c2:
            st.download_button("📁 Download DOCX", generate_docx_executive(cv_f, st.session_state.active_lang),
                               file_name=f"{f_name}.docx", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)